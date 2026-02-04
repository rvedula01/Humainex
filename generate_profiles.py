import os
import json
import random
import re
from datetime import datetime, timedelta
from faker import Faker
from openai import OpenAI
import httpx
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from config import get_api_key

# print(os.environ.get('HTTP_PROXY'))
# print(os.environ.get('HTTPS_PROXY'))
# print(os.environ.get('http_proxy'))
# print(os.environ.get('https_proxy'))

# Load environment variables and set up API key
load_dotenv()
try:
    OpenAI.api_key = get_api_key()
except Exception as e:
    print(f"Error initializing API key: {e}")
    raise

# Create a custom HTTP client without proxy settings
http_client = httpx.Client(
    timeout=60.0,
    verify=True,
    follow_redirects=True
)

# Initialize OpenAI client with custom HTTP client
client = OpenAI(
    api_key=OpenAI.api_key,
    http_client=http_client
)

# Initialize Faker
fake = Faker()

# Role definitions with required skills
ROLES = {
    'Data Scientist': {
        'skills': ['Python', 'Machine Learning', 'Statistics', 'Data Analysis', 'SQL', 'Deep Learning', 'Data Visualization', 'NLP'],
        'experience_range': (3, 5),
        'education': [
            'MS in Data Science',
            'PhD in Computer Science',
            'MS in Statistics',
            'Masters in Artificial Intelligence'
        ]
    },
    'Data Engineer': {
        'skills': ['Python', 'SQL', 'ETL', 'Data Warehousing', 'Spark', 'Hadoop', 'AWS/GCP'],
        'experience_range': (3, 5),
        'education': [
            'BS in Computer Science',
            'MS in Computer Engineering',
            'BS in Information Technology'
        ]
    },
    'Data Analyst': {
        'skills': ['SQL', 'Excel', 'Tableau', 'Data Visualization', 'Python', 'Statistics'],
        'experience_range': (3, 5),
        'education': [
            'BS in Statistics',
            'BS in Economics',
            'BS in Business Analytics'
        ]
    },
    'Product Analyst': {
        'skills': ['SQL', 'A/B Testing', 'Product Metrics', 'Data Visualization', 'Excel', 'Python'],
        'experience_range': (3, 5),
        'education': [
            'BS in Business',
            'BS in Economics',
            'MS in Business Analytics'
        ]
    },
    'Python Developer': {
        'skills': ['Python', 'Django/Flask', 'REST APIs', 'Git', 'Docker', 'SQL', 'Testing'],
        'experience_range': (3, 5),
        'education': [
            'BS in Computer Science',
            'BS in Software Engineering',
            'MS in Computer Science'
        ]
    }
}

def generate_candidate_profile(role):
    """Generate a candidate profile with relevant skills and experience."""
    role_info = ROLES[role]
    experience = random.randint(*role_info['experience_range'])
    
    profile = {
        'name': fake.name(),
        'email': fake.email(),
        'phone': fake.phone_number(),
        'location': fake.city() + ', ' + fake.country(),
        'role': role,
        'experience_years': experience,
        'education': random.choice(role_info['education']),
        'university': fake.company()+ " University",
        'skills': random.sample(role_info['skills'], k=min(6, len(role_info['skills']))),
        'summary': f"Experienced {role} with {experience} years of experience in {', '.join(role_info['skills'][:3])}."
    }
    
    # Add some random skills (20% chance for each)
    all_skills = set()
    for r in ROLES.values():
        all_skills.update(r['skills'])
    
    additional_skills = random.sample(
        list(all_skills - set(profile['skills'])),
        k=min(3, len(all_skills - set(profile['skills'])))
    )
    profile['skills'].extend(additional_skills)
    
    return profile

def generate_resume_content(profile):
    """Generate detailed resume content using OpenAI."""
    # Set 2025 as the current application year
    current_year = 2025
    # Calculate graduation year (assume 0-2 years between graduation and first job)
    graduation_year = current_year - profile['experience_years'] - random.randint(0, 2)
    
    prompt = f"""
    Create a professional one-page resume for a {profile['role']} candidate with the following details:
    
    Name: {profile['name']}
    Experience: {profile['experience_years']} years
    Education: {profile['education']} from {profile['university']}
    Skills: {', '.join(profile['skills'])}
    
    Formatting Instructions:
    - Use actual section headers (not markdown) for section titles
    - Do not use ** or any markdown syntax for bold/formatting
    - Use proper line breaks between sections
    - Use bullet points for lists and achievements
    - Use Aug 2025 as the current date for all date references
    - For education, use graduation year: {graduation_year}
    - For work experience, ensure dates align with the total {profile['experience_years']} years of experience
    - All dates should be between {graduation_year} and Aug 2025
    - For current positions, use 'Present' instead of Aug 2025 (e.g., '2023 - Present')
    
    Include the following sections:
    
    Professional Summary
    [3-4 sentences highlighting key achievements]
    
    Technical Skills
    [Categorized by area, no bullet points needed]
    
    Professional Experience
    [Company Name], [Location] - [Job Title] (YYYY-YYYY or YYYY-Present)
    • [Achievement 1]
    • [Achievement 2]
    • [Achievement 3]
    
    [Previous Company], [Location] - [Job Title] (YYYY-YYYY)
    • [Achievement 1]
    • [Achievement 2]
    
    Education
    {profile['education']}
    {profile['university']}, {graduation_year}
    
    Certifications
    • [Relevant Certification] - {random.randint(graduation_year + 1, current_year)}
    • [Relevant Certification] - {random.randint(graduation_year + 1, current_year)}
    
    Do not include any additional content after the Certifications section.
    Make the content realistic and tailored for a {profile['role']} role.
    """
    
    try:        
        # Create a custom HTTP client with no proxies
        http_client = httpx.Client(
            timeout=60.0,  # 60 second timeout
            verify=True,   # Verify SSL certificates
            follow_redirects=True
        )
        
        # Initialize the OpenAI client with the custom HTTP client
        client = OpenAI(
            api_key=openai.api_key,
            http_client=http_client
        )
        
        # Make the API call
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a professional resume writer creating a one-page resume."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=1500,
            temperature=0.7
        )
        
        # Clean up the HTTP client
        http_client.close()
        
        return response.choices[0].message.content.strip()
    except Exception as e:
        print(f"Error generating resume content: {e}")
        if 'http_client' in locals():
            http_client.close()
        return ""

def generate_job_description():
    """Generate a job description for a Data Scientist with 4 years of experience."""
    try:        
        # Create a custom HTTP client
        http_client = httpx.Client(
            timeout=60.0,
            verify=True,
            follow_redirects=True
        )
        
        # Initialize the OpenAI client
        client = OpenAI(
            api_key=openai.api_key,
            http_client=http_client
        )
        
        prompt = """
        Create a 2-pages detailed job description for a Data Scientist position requiring 4 years of experience.
        
        Include the following sections:
        1. Job Title: Data Scientist (4+ years experience)
        2. Location: [Specify or mention 'Remote']
        3. Job Type: [Full-time/Contract]
        4. About the Role:
           - Brief overview of the position
           - Team and company context
        
        5. Key Responsibilities:
           - 6-8 bullet points of core responsibilities
           - Focus on data analysis, modeling, and business impact
           
        6. Requirements:
           - 4+ years of relevant experience
           - Technical skills (Python, SQL, ML frameworks, etc.)
           - Educational qualifications
           - Any specific domain knowledge
           
        7. Nice-to-Have Skills:
           - Additional technologies or experience
           - Soft skills
           
        8. What We Offer:
           - Company benefits
           - Growth opportunities
           - Work culture
        
        Make it professional, detailed, and realistic for a mid-level position.
        """
        
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a professional technical recruiter creating a job description."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=1500,
            temperature=0.7
        )
        
        # Clean up
        http_client.close()
        
        # Create output directory if it doesn't exist
        os.makedirs('job_descriptions', exist_ok=True)
        
        # Create a new Document
        doc = Document()
        
        # Add title
        title = doc.add_heading('Job Description: Data Scientist (4+ years experience)', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add content with formatting
        content = response.choices[0].message.content
        
        # Split content into paragraphs and add to document
        for paragraph in content.split('\n\n'):
            if paragraph.strip() == '':
                continue
                
            # Check if this is a section header
            if ':' in paragraph and len(paragraph) < 50:  # Simple heuristic for header detection
                p = doc.add_heading(paragraph.strip(' :'), level=2)
            else:
                # Handle bullet points
                if paragraph.strip().startswith('•') or paragraph.strip().startswith('-'):
                    for line in paragraph.split('\n'):
                        if line.strip().startswith(('•', '-')):
                            p = doc.add_paragraph(style='List Bullet')
                            p.add_run(line.lstrip('•- ').strip())
                else:
                    doc.add_paragraph(paragraph.strip())
        
        # Save the document
        file_path = os.path.join('job_descriptions', 'Data_Scientist_JD.docx')
        doc.save(file_path)
        
        print(f"Job description saved to {file_path}")
        return file_path
        
    except Exception as e:
        print(f"Error generating job description: {e}")
        if 'http_client' in locals():
            http_client.close()
        return None

def create_word_document(profile, content, output_dir='resumes'):
    """Create a Word document for the candidate's resume."""
    # Create output directory if it doesn't exist
    os.makedirs(output_dir, exist_ok=True)
    
    # Create a new Document
    doc = Document()
    
    # Add title
    title = doc.add_heading(profile['name'], level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add contact information
    contact = doc.add_paragraph()
    contact.add_run(f"{profile['email']} | {profile['phone']} | {profile['location']}")
    contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add a line break
    doc.add_paragraph()
    
    # Add the generated content
    for line in content.split('\n'):
        if line.strip() == '':
            doc.add_paragraph()
            continue
            
        if line.strip().endswith(':'):
            # This is a section header
            doc.add_heading(line.strip(' :'), level=2)
        else:
            # This is regular content
            p = doc.add_paragraph()
            if line.strip().startswith('•'):
                # This is a bullet point
                p.add_run('• ' + line[1:].strip())
            else:
                p.add_run(line.strip())
    
    # Save the document
    filename = f"{profile['name'].replace(' ', '_')}_Resume.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    return filepath

def main():
    print("Starting resume processing...")
    
    # Uncomment the following section to generate new profiles
    
    print("Generating candidate profiles...")
    
    # Create 3 data scientist profiles
    # profiles = [generate_candidate_profile('Data Scientist') for _ in range(3)]
    
    # Create 7 profiles for other roles
    # other_roles = [r for r in ROLES.keys() if r != 'Data Scientist']
    # profiles += [generate_candidate_profile(random.choice(other_roles)) for _ in range(7)]
    
    # Shuffle the profiles
    # random.shuffle(profiles)
    
    # Generate resumes for each profile
    # for i, profile in enumerate(profiles, 1):
    #     print(f"\nGenerating resume for candidate {i}: {profile['name']} ({profile['role']})")
        
    #     # Generate resume content
    #     resume_content = generate_resume_content(profile)
        
    #     # Create Word document
    #     if resume_content:
    #         filepath = create_word_document(profile, resume_content)
    #         print(f"Created resume: {filepath}")
    #     else:
    #         print(f"Failed to generate resume for {profile['name']}")
    
    # print("\nResume generation complete!")
    
    # Generate job description
    # print("\nGenerating job description for Data Scientist (4+ years experience)...")
    # jd_path = generate_job_description()
    # if jd_path:
    #     print(f"Job description saved to: {os.path.abspath(jd_path)}")
    
    
    # Process existing resumes and export to Excel
    #     if resume_content:
    #         filepath = create_word_document(profile, resume_content)
    #         print(f"Created resume: {filepath}")
    #     else:
    #         print(f"Failed to generate resume for {profile['name']}")
    
    # print("\nResume generation complete!")
    
    # Generate job description
    # print("\nGenerating job description for Data Scientist (4+ years experience)...")
    # jd_path = generate_job_description()
    # if jd_path:
    #     print(f"Job description saved to: {os.path.abspath(jd_path)}")

def extract_resume_with_gpt(resume_text: str) -> dict:
    """Extract structured information from resume text using GPT-4o-mini."""
    try:
        system_prompt = """Extract resume details including name, contact info, experience, education, skills, certifications, and Tags. 
        Return as JSON with keys: name, email, phone, location, summary, experience, education, skills, certifications, tags.
        Experience and education should be lists of dicts with relevant fields."""
        
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"Extract information from this resume:\n\n{resume_text}"}
            ],
            temperature=0.3,
            max_tokens=2000
        )
        
        return json.loads(response.choices[0].message['content'])
    except Exception as e:
        print(f"Error extracting data with GPT: {str(e)}")
        return {}

def extract_text_from_docx(docx_path: str) -> str:
    """Extract text from a docx file."""
    doc = Document(docx_path)
    return '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])

def format_experience(experience: list) -> str:
    """Format experience list into a readable string."""
    if not experience:
        return ""
    return '\n\n'.join(
        f"{exp.get('position', '')} at {exp.get('company', '')} ({exp.get('duration', 'N/A')})\n" +
        '\n'.join(f"- {resp}" for resp in exp.get('responsibilities', []))
        for exp in experience
    )

def format_education(education: list) -> str:
    """Format education list into a readable string."""
    if not education:
        return ""
    return '\n'.join(
        f"{edu.get('degree', '')} from {edu.get('institution', '')} ({edu.get('year', 'N/A')})"
        for edu in education
    )

def export_resumes_to_excel(resumes_dir: str = 'resumes', output_file: str = 'candidate_resumes.xlsx') -> str:
    """
    Extract sections from all resumes in the directory using GPT-4 and save to an Excel file.
    
    Args:
        resumes_dir: Directory containing resume files
        output_file: Path to save the Excel file
        
    Returns:
        str: Path to the generated Excel file
    """
    import pandas as pd
    import os
    from tqdm import tqdm
    
    # Define output columns
    columns = [
        'Name', 'Email', 'Phone', 'Location', 'Summary', 'Experience',
        'Education', 'Skills', 'Certifications', 'Tags', 'Source File'
    ]
    
    # Initialize DataFrame with columns
    df = pd.DataFrame(columns=columns)
    
    # Get all docx files in the directory
    resume_files = [f for f in os.listdir(resumes_dir) if f.lower().endswith('.docx')]
    
    if not resume_files:
        print(f"No resume files found in {resumes_dir}")
        return ""
    
    print(f"Processing {len(resume_files)} resumes with GPT-4...")
    
    # Process each resume
    for resume_file in tqdm(resume_files, desc="Processing resumes"):
        try:
            # Get full path to resume
            docx_path = os.path.join(resumes_dir, resume_file)
            
            # Extract text from docx
            resume_text = extract_text_from_docx(docx_path)
            
            # Extract information using GPT-4
            extracted_data = extract_resume_with_gpt(resume_text)
            
            if not extracted_data:
                print(f"Warning: Could not extract data from {resume_file}")
                continue
            
            # Format the data for Excel
            candidate_data = {
                'Name': extracted_data.get('name', '').strip() or os.path.splitext(resume_file)[0].replace('_', ' '),
                'Email': extracted_data.get('email', '').strip(),
                'Phone': extracted_data.get('phone', '').strip(),
                'Location': extracted_data.get('location', '').strip(),
                'Summary': extracted_data.get('summary', '').strip(),
                'Experience': format_experience(extracted_data.get('experience', [])),
                'Education': format_education(extracted_data.get('education', [])),
                'Skills': ', '.join(extracted_data.get('skills', [])),
                'Certifications': ', '.join(extracted_data.get('certifications', [])),
                'Tags': ', '.join(extracted_data.get('tags', [])),
                'Source File': resume_file
            }
            
            # Add to DataFrame
            df = pd.concat([df, pd.DataFrame([candidate_data])], ignore_index=True)
            
        except Exception as e:
            print(f"Error processing {resume_file}: {str(e)}")
    
    if df.empty:
        print("No valid resume data to export.")
        return ""
    
    # Ensure output directory exists
    os.makedirs(os.path.dirname(os.path.abspath(output_file)) or '.', exist_ok=True)
    
    # Save to Excel with auto-adjusted column widths
    try:
        with pd.ExcelWriter(output_file, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Resumes')
            
            # Auto-adjust column widths
            worksheet = writer.sheets['Resumes']
            for idx, col in enumerate(df.columns):
                max_length = max((
                    df[col].astype(str).apply(len).max(),
                    len(str(col))
                ))
                # Add a little extra space
                worksheet.column_dimensions[chr(65 + idx)].width = min(max_length + 2, 50)
        
        print(f"\nSuccessfully exported {len(df)} resumes to {os.path.abspath(output_file)}")
        return output_file
        
    except Exception as e:
        print(f"Error saving to Excel: {str(e)}")
        # Fallback to simple save if there's an error with formatting
        df.to_excel(output_file, index=False, engine='openpyxl')
        print(f"Exported with basic formatting to {os.path.abspath(output_file)}")
        return output_file

if __name__ == "__main__":
    # Process existing resumes and export to Excel
    print("\nExporting resume data to Excel...")
    excel_file = export_resumes_to_excel()
    if excel_file:
        print(f"Resume data exported to {excel_file}")
    
    # Uncomment the following to generate new profiles and job descriptions
    
    # Generate job description
    # print("\nGenerating job description for Data Scientist (3+ years experience)...")
    # jd_path = generate_job_description()
    # if jd_path:
    #     print(f"Job description saved to: {os.path.abspath(jd_path)}")
    
    # Call the main function to generate new profiles
    main()
